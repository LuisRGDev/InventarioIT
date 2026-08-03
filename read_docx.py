from docx import Document
import sys

doc = Document("Formato para llenar Responsiva.docx")

print("=== PARRAFOS ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        runs_detail = []
        for r in p.runs:
            bold = "B" if r.bold else ""
            italic = "I" if r.italic else ""
            size = r.font.size.pt if r.font.size else "?"
            runs_detail.append(f"[{bold}{italic} sz={size}] '{r.text}'")
        print(f"P{i} ({p.alignment}): [{p.style.name}] {p.text}")
        for rd in runs_detail:
            print(f"    Run: {rd}")

print()
print("=== TABLAS ===")
for t_idx, table in enumerate(doc.tables):
    print(f"--- Tabla {t_idx} ({len(table.rows)} filas x {len(table.columns)} cols) ---")
    for r_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  Fila {r_idx}: {cells}")

print()
print("=== HEADERS/FOOTERS ===")
for section in doc.sections:
    header = section.header
    if header and not header.is_linked_to_previous:
        for p in header.paragraphs:
            if p.text.strip():
                print(f"Header: {p.text}")
        for table in header.tables:
            for r_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                print(f"  Header Table Fila {r_idx}: {cells}")
    footer = section.footer
    if footer and not footer.is_linked_to_previous:
        for p in footer.paragraphs:
            if p.text.strip():
                print(f"Footer: {p.text}")

print()
print("=== IMAGENES ===")
from docx.opc.constants import RELATIONSHIP_TYPE as RT
rels = doc.part.rels
img_count = 0
for rel in rels.values():
    if "image" in rel.reltype:
        img_count += 1
        print(f"  Image: {rel.target_ref}")
print(f"Total images in body: {img_count}")

# Check header images
for section in doc.sections:
    header = section.header
    if header:
        for rel in header.part.rels.values():
            if "image" in rel.reltype:
                print(f"  Header image: {rel.target_ref}")
